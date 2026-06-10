"""Deterministic python-docx resume formatter."""

from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from functools import lru_cache
from pathlib import Path
from typing import Any

import mammoth
from pypdf import PdfReader

import frontmatter
from ai import ai_output_char_count, normalize_ai_output
from docx import Document
from scoring import (
    DEFAULT_OVERFLOW_WARNING_MIN_COMPOSITE,
    SelectionPolicy,
    expand_selection_by_highest_score,
    high_quality_omitted_items,
    selection_with_all_items,
    trim_selection_by_lowest_score,
)
from selection import (
    DEFAULT_MIN_PROJECT_BULLETS,
    DEFAULT_MIN_PROJECT_ENTRIES,
    apply_content_selection,
    default_selection,
    selection_trim_exhausted,
    static_content_stats,
    trim_selection_one_step,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

FONT_NAME = "Calibri"
FONT_BODY = Pt(10)
FONT_NAME_SIZE = Pt(14)
FONT_SECTION = Pt(11)
MARGIN_TB = Inches(0.5)
MARGIN_LR = Inches(0.5)
SECTION_SPACING_BEFORE = Pt(7)
SECTION_SPACING_AFTER = Pt(3)
BULLET_TEXT_INDENT = Inches(0.25)

ACCENT_COLOR = RGBColor(0x1A, 0x56, 0x8C)
RULE_COLOR_HEX = "1A568C"
META_COLOR = RGBColor(0x55, 0x55, 0x55)
LINK_COLOR = RGBColor(0x1A, 0x56, 0x8C)

DEFAULT_RESUME_SECTIONS = [
    "summary",
    "skills",
    "experience",
    "education",
    "projects",
    "certifications",
]
VALID_RESUME_SECTIONS = frozenset(DEFAULT_RESUME_SECTIONS)
STATIC_PREVIEW_SECTIONS = frozenset({"summary", "skills"})


def resolve_resume_sections(sections: list[str] | None) -> list[str]:
    """Validate and normalize section ids; fall back to default when unset or empty."""
    if not sections:
        return list(DEFAULT_RESUME_SECTIONS)

    unknown = [section for section in sections if section not in VALID_RESUME_SECTIONS]
    if unknown:
        raise ValueError(
            f"Unknown resume section(s): {', '.join(unknown)}. "
            f"Valid ids: {', '.join(sorted(VALID_RESUME_SECTIONS))}"
        )

    seen: set[str] = set()
    resolved: list[str] = []
    for section in sections:
        if section not in seen:
            seen.add(section)
            resolved.append(section)
    return resolved if resolved else list(DEFAULT_RESUME_SECTIONS)


def ai_section_flags(sections: list[str] | None) -> tuple[bool, bool]:
    """Return (include_summary, include_skills) from resolved resume section config."""
    resolved = resolve_resume_sections(sections)
    return "summary" in resolved, "skills" in resolved


def static_preview_sections(sections: list[str] | None) -> list[str]:
    """Sections for no-LLM background preview (drops summary/skills)."""
    return [s for s in resolve_resume_sections(sections) if s not in STATIC_PREVIEW_SECTIONS]


def load_background(path: Path) -> dict[str, Any]:
    """Load and validate YAML frontmatter from background.md."""
    if not path.exists():
        raise FileNotFoundError(f"Background file not found: {path}")
    post = frontmatter.load(path)
    return validate_background(post.metadata)


def validate_background(metadata: dict[str, Any]) -> dict[str, Any]:
    """Validate frontmatter schema; return metadata unchanged if valid."""
    if not isinstance(metadata, dict):
        raise ValueError("Background frontmatter must be a YAML mapping")

    header = metadata.get("header")
    if not isinstance(header, dict):
        raise ValueError("background.md: 'header' must be a mapping")
    for key in ("name", "email"):
        if not header.get(key):
            raise ValueError(f"background.md: header.{key} is required")

    experience = metadata.get("experience")
    if not isinstance(experience, list) or not experience:
        raise ValueError("background.md: 'experience' must be a non-empty list")

    for i, job in enumerate(experience):
        if not isinstance(job, dict):
            raise ValueError(f"background.md: experience[{i}] must be a mapping")
        for key in ("company", "title", "start", "end", "bullets"):
            if key not in job:
                raise ValueError(f"background.md: experience[{i}].{key} is required")
        if not isinstance(job["bullets"], list) or not job["bullets"]:
            raise ValueError(f"background.md: experience[{i}].bullets must be non-empty")

    for section, required_keys in (
        ("education", ("institution", "degree", "graduation")),
        ("certifications", ("name",)),
        ("projects", ("name", "bullets")),
    ):
        items = metadata.get(section, [])
        if items is None:
            metadata[section] = []
            continue
        if not isinstance(items, list):
            raise ValueError(f"background.md: '{section}' must be a list")
        for i, item in enumerate(items):
            if not isinstance(item, dict):
                raise ValueError(f"background.md: {section}[{i}] must be a mapping")
            for key in required_keys:
                if key not in item:
                    raise ValueError(f"background.md: {section}[{i}].{key} is required")

    links = header.get("links", [])
    if links is not None and not isinstance(links, list):
        raise ValueError("background.md: header.links must be a list")

    return metadata


def _render_summary_section(
    doc: Document,
    data: dict[str, Any],
    normalized: dict[str, Any],
) -> None:
    _add_section_heading(doc, "Professional Summary")
    _add_body_paragraph(doc, normalized["summary"])


def _render_skills_section(
    doc: Document,
    data: dict[str, Any],
    normalized: dict[str, Any],
) -> None:
    categories = normalized["skill_categories"]
    if not categories:
        return
    _add_section_heading(doc, "Skills")
    _add_skill_categories(doc, categories)


def _render_experience_section(
    doc: Document,
    data: dict[str, Any],
    normalized: dict[str, Any],
) -> None:
    experience = data.get("experience", [])
    if experience:
        _add_section_heading(doc, "Experience")
        for i, job in enumerate(experience):
            _add_experience_entry(doc, job, is_first=(i == 0))


def _render_education_section(
    doc: Document,
    data: dict[str, Any],
    normalized: dict[str, Any],
) -> None:
    education = data.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        for entry in education:
            _add_education_entry(doc, entry)


def _render_projects_section(
    doc: Document,
    data: dict[str, Any],
    normalized: dict[str, Any],
) -> None:
    projects = data.get("projects", [])
    if projects:
        _add_section_heading(doc, "Projects")
        for project in projects:
            _add_project_entry(doc, project)


def _render_certifications_section(
    doc: Document,
    data: dict[str, Any],
    normalized: dict[str, Any],
) -> None:
    certifications = data.get("certifications", [])
    if certifications:
        _add_section_heading(doc, "Certifications")
        for cert in certifications:
            _add_certification_entry(doc, cert)


_SECTION_RENDERERS: dict[str, Any] = {
    "summary": _render_summary_section,
    "skills": _render_skills_section,
    "experience": _render_experience_section,
    "education": _render_education_section,
    "projects": _render_projects_section,
    "certifications": _render_certifications_section,
}


def build_resume(
    data: dict[str, Any],
    ai_output: dict[str, Any],
    *,
    sections: list[str] | None = None,
    content_selection: dict[str, Any] | None = None,
    max_certifications: int | None = None,
) -> bytes:
    """Merge structured background + AI fields, return DOCX bytes."""
    render_data = (
        apply_content_selection(data, content_selection)
        if content_selection is not None
        else data
    )
    if max_certifications is not None:
        render_data = {
            **render_data,
            "certifications": render_data.get("certifications", [])[:max_certifications],
        }
    resolved_sections = resolve_resume_sections(sections)
    include_summary, include_skills = ai_section_flags(resolved_sections)
    normalized = normalize_ai_output(
        ai_output,
        include_summary=include_summary,
        include_skills=include_skills,
    )

    doc = Document()
    _set_document_margins(doc)
    _set_default_font(doc)

    _add_header(doc, render_data["header"])
    for section_id in resolved_sections:
        _SECTION_RENDERERS[section_id](doc, render_data, normalized)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_document_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = MARGIN_TB
        section.bottom_margin = MARGIN_TB
        section.left_margin = MARGIN_LR
        section.right_margin = MARGIN_LR


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_BODY
    rfonts = style.element.rPr.rFonts if style.element.rPr is not None else None
    if rfonts is None:
        rpr = style.element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def _format_run(run, *, bold: bool = False, size: Pt | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = size or FONT_BODY
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_paragraph_border_bottom(paragraph, color_hex: str = RULE_COLOR_HEX) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SECTION_SPACING_BEFORE
    p.paragraph_format.space_after = SECTION_SPACING_AFTER
    run = p.add_run(title.upper())
    _format_run(run, bold=True, size=FONT_SECTION, color=ACCENT_COLOR)
    _add_paragraph_border_bottom(p)


def _add_hyperlink(paragraph, url: str, display_text: str) -> None:
    """Insert a genuine clickable hyperlink run into paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), RULE_COLOR_HEX)
    r_pr.append(color_el)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    r_pr.append(sz)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_pr.append(r_fonts)

    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = display_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _format_run(run)


def _add_header(doc: Document, header: dict[str, Any]) -> None:
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(header["name"])
    _format_run(name_run, bold=True, size=FONT_NAME_SIZE)

    email = header.get("email")
    contact_parts = [
        header.get("phone"),
        header.get("location"),
    ]
    contact_values = [p for p in contact_parts if p]
    links = header.get("links") or []
    valid_links = [
        link for link in links if link.get("label") and link.get("url")
    ]
    if email or contact_values or valid_links:
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.paragraph_format.space_after = Pt(8)

        first = True
        if email:
            _add_hyperlink(meta_p, f"mailto:{email}", email)
            first = False

        for part in contact_values:
            if not first:
                sep_run = meta_p.add_run(" | ")
                _format_run(sep_run, size=Pt(10), color=META_COLOR)
            part_run = meta_p.add_run(part)
            _format_run(part_run, size=Pt(10), color=META_COLOR)
            first = False

        for link in valid_links:
            if not first:
                sep_run = meta_p.add_run(" | ")
                _format_run(sep_run, size=Pt(10), color=META_COLOR)
            _add_hyperlink(meta_p, link["url"], link["label"])
            first = False


def _configure_bullet_paragraph(p) -> None:
    """Hanging indent so wrapped lines align with the first line's text column."""
    fmt = p.paragraph_format
    fmt.left_indent = BULLET_TEXT_INDENT
    fmt.first_line_indent = -BULLET_TEXT_INDENT
    fmt.tab_stops.add_tab_stop(BULLET_TEXT_INDENT)


def _add_skill_categories(doc: Document, skill_categories: list[dict[str, Any]]) -> None:
    """Render one bullet per skill category with comma-joined skills."""
    if not skill_categories:
        return

    for cat in skill_categories:
        skills_text = ", ".join(cat["skills"])
        p = doc.add_paragraph()
        _configure_bullet_paragraph(p)
        p.paragraph_format.space_after = Pt(2)
        bullet_run = p.add_run("▪")
        _format_run(bullet_run, color=ACCENT_COLOR)
        p.add_run("\t")
        name_run = p.add_run(f"{cat['name']}: ")
        _format_run(name_run, bold=True)
        text_run = p.add_run(skills_text)
        _format_run(text_run)


def _normalize_date(value: str) -> str:
    if value.lower() == "present":
        return "Present"
    return value


def _format_date_range(start: str, end: str) -> str:
    return f"{_normalize_date(start)} – {_normalize_date(end)}"


def _add_experience_entry(
    doc: Document, job: dict[str, Any], *, is_first: bool = False
) -> None:
    title_line = doc.add_paragraph()
    title_line.paragraph_format.space_before = Pt(0 if is_first else 5)
    title_line.paragraph_format.space_after = Pt(0)

    role_run = title_line.add_run(job["title"])
    _format_run(role_run, bold=True)

    sep = title_line.add_run("  ·  ")
    _format_run(sep, color=META_COLOR)

    company_run = title_line.add_run(job["company"])
    _format_run(company_run, color=META_COLOR)

    if job.get("location"):
        loc_sep = title_line.add_run("  ·  ")
        _format_run(loc_sep, color=META_COLOR)
        loc_run = title_line.add_run(job["location"])
        _format_run(loc_run, color=META_COLOR, size=Pt(9.5))

    tab_run = title_line.add_run("\t" + _format_date_range(job["start"], job["end"]))
    _format_run(tab_run, color=META_COLOR, size=Pt(9.5))

    p_pr = title_line._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9072")
    tabs.append(tab)
    p_pr.append(tabs)

    for bullet in job["bullets"]:
        _add_bullet(doc, bullet)


def _add_education_entry(doc: Document, entry: dict[str, Any]) -> None:
    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(2)
    line.paragraph_format.space_after = Pt(4)
    inst_run = line.add_run(entry["institution"])
    _format_run(inst_run, bold=True)
    if entry.get("location"):
        loc_run = line.add_run(f" — {entry['location']}")
        _format_run(loc_run)
    degree_run = line.add_run(f" | {entry['degree']}")
    _format_run(degree_run)
    if entry.get("graduation"):
        grad_run = line.add_run(f" | {entry['graduation']}")
        _format_run(grad_run)


def _add_certification_entry(doc: Document, cert: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    name_run = p.add_run(cert["name"])
    _format_run(name_run, bold=True)
    extras = []
    if cert.get("issuer"):
        extras.append(cert["issuer"])
    if cert.get("date"):
        extras.append(str(cert["date"]))
    if extras:
        extra_run = p.add_run(f" — {', '.join(extras)}")
        _format_run(extra_run)


def _add_project_entry(doc: Document, project: dict[str, Any]) -> None:
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(0)

    name_run = title_p.add_run(project["name"])
    _format_run(name_run, bold=True)

    if project.get("url"):
        url_sep = title_p.add_run("  —  ")
        _format_run(url_sep, size=Pt(9.5), color=META_COLOR)
        _add_hyperlink(title_p, project["url"], project["url"])

    tech = project.get("tech") or project.get("stack")
    if tech:
        tech_run = title_p.add_run(f"  —  {tech}")
        _format_run(tech_run, size=Pt(9.5), color=META_COLOR)

    for bullet in project.get("bullets", []):
        _add_bullet(doc, bullet)


def _prevent_orphan_word(text: str) -> str:
    """Join last 2 words with non-breaking spaces to avoid a lone short word wrapping."""
    words = text.split()
    if len(words) < 3:
        return text
    head, tail = words[:-2], words[-2:]
    return " ".join(head) + "\u00a0" + "\u00a0".join(tail)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _configure_bullet_paragraph(p)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:widowControl"))
    bullet_run = p.add_run("▪")
    _format_run(bullet_run, color=ACCENT_COLOR)
    p.add_run("\t")
    text_run = p.add_run(_prevent_orphan_word(text))
    _format_run(text_run)


PREVIEW_HTML_STYLE = """
<style>
  html, body {
    background: #ffffff;
    font-family: Calibri, "Segoe UI", sans-serif;
    font-size: 10pt;
    line-height: 1.35;
    color: #222;
    max-width: 8.5in;
    margin: 0 auto;
    padding: 0.5in 0.6in;
  }
  p { margin: 0.25em 0; }
  a { color: #1A568C; }
  strong { color: #1A568C; }
</style>
"""


def resume_filename(candidate_name: str, ext: str = "docx") -> str:
    """Build a stable resume filename from the candidate name in background.md."""
    safe = re.sub(r"[^\w\-]+", "_", candidate_name.strip()).strip("_")
    normalized_ext = ext.lstrip(".") or "docx"
    base = f"{safe}_Resume" if safe else "resume"
    return f"{base}.{normalized_ext}"


def save_resume_to_disk(file_bytes: bytes, output_dir: Path, filename: str) -> Path:
    """Write resume bytes to output_dir, overwriting any existing file with the same name."""
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = Path(filename).name
    if not safe_name or safe_name in (".", ".."):
        safe_name = "resume.docx"
    path = output_dir / safe_name
    path.write_bytes(file_bytes)
    return path


@lru_cache(maxsize=1)
def libreoffice_available() -> bool:
    """Return True when LibreOffice headless conversion is available."""
    return shutil.which("soffice") is not None


def docx_to_html(docx_bytes: bytes) -> str:
    """Convert DOCX bytes to styled HTML for in-app preview."""
    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    body = result.value
    return f"<!DOCTYPE html><html><head>{PREVIEW_HTML_STYLE}</head><body>{body}</body></html>"


def docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    """Convert DOCX bytes to PDF via LibreOffice headless, or None if unavailable."""
    if not libreoffice_available():
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(docx_bytes)
        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_path),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                timeout=120,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            return None

        pdf_path = tmp_path / "resume.pdf"
        if not pdf_path.exists():
            return None
        return pdf_path.read_bytes()


def pdf_page_count(pdf_bytes: bytes) -> int:
    """Return the number of pages in a PDF."""
    return len(PdfReader(io.BytesIO(pdf_bytes)).pages)


def trim_summary_one_step(
    ai_output: dict[str, Any],
    *,
    sections: list[str] | None = None,
) -> dict[str, Any]:
    """Remove one summary sentence to reduce rendered page count. Skills are never trimmed."""
    include_summary, include_skills = ai_section_flags(sections)
    normalized = normalize_ai_output(
        ai_output,
        include_summary=include_summary,
        include_skills=include_skills,
    )
    if not include_summary:
        return normalized

    summary = normalized["summary"]
    categories = [
        {"name": cat["name"], "skills": list(cat["skills"])}
        for cat in normalized["skill_categories"]
    ]
    sentences = re.split(r"(?<=[.!?])\s+", summary)
    if len(sentences) > 1:
        sentences.pop()
        return {"summary": " ".join(sentences).strip(), "skill_categories": categories}
    return normalized


# Backward-compatible alias for tests migrating from skill-trimming behavior.
trim_ai_output_one_step = trim_summary_one_step


def fit_resume_to_page_budget(
    background_data: dict[str, Any],
    ai_output: dict[str, Any],
    *,
    max_pages: int = 2,
    max_chars: int | None = None,
    sections: list[str] | None = None,
    content_selection: dict[str, Any] | None = None,
    max_certifications: int | None = None,
    min_project_entries: int = DEFAULT_MIN_PROJECT_ENTRIES,
    min_project_bullets: int = DEFAULT_MIN_PROJECT_BULLETS,
    auto_fill_page_budget: bool = True,
    overflow_warning_min_composite: float = DEFAULT_OVERFLOW_WARNING_MIN_COMPOSITE,
) -> tuple[dict[str, Any], dict[str, Any], bytes, bytes | None, int | None, dict[str, Any]]:
    """
    Expand or trim selection (and summary) until the rendered PDF fits max_pages.

    Returns (fitted_ai_output, fitted_selection, docx_bytes, pdf_bytes, page_count, diagnostics).
    page_count is None when LibreOffice is unavailable.
    """
    include_summary, include_skills = ai_section_flags(sections)
    current = normalize_ai_output(
        ai_output,
        include_summary=include_summary,
        include_skills=include_skills,
    )
    if content_selection is None:
        current_selection = default_selection(
            background_data,
            min_project_entries=min_project_entries,
            min_project_bullets=min_project_bullets,
        )
    else:
        current_selection = {
            "experience_selections": [
                {"role_index": e["role_index"], "bullet_indices": list(e["bullet_indices"])}
                for e in content_selection.get("experience_selections", [])
            ],
            "project_selections": [
                {"project_index": p["project_index"], "bullet_indices": list(p["bullet_indices"])}
                for p in content_selection.get("project_selections", [])
            ],
            "education_indices": list(content_selection.get("education_indices", [])),
        }
        for key in ("content_ratings", "content_composites", "trim_log", "expand_log"):
            if key in content_selection:
                current_selection[key] = deepcopy(content_selection[key])

    def _render_and_measure(
        selection: dict[str, Any],
        ai: dict[str, Any],
    ) -> tuple[bytes, bytes | None, int | None]:
        docx = build_resume(
            background_data,
            ai,
            sections=sections,
            content_selection=selection,
            max_certifications=max_certifications,
        )
        pdf = docx_to_pdf(docx)
        pages = pdf_page_count(pdf) if pdf else None
        return docx, pdf, pages

    docx_bytes, pdf_bytes, page_count = _render_and_measure(current_selection, current)
    trim_stalled = False
    if pdf_bytes is None:
        diagnostics = _page_fit_diagnostics(
            background_data,
            current,
            current_selection,
            page_count=None,
            max_pages=max_pages,
            max_chars=max_chars,
            include_summary=include_summary,
            include_skills=include_skills,
            trim_stalled=False,
            overflow_warning_min_composite=overflow_warning_min_composite,
        )
        return current, current_selection, docx_bytes, None, None, diagnostics

    trim_policy = SelectionPolicy(
        min_project_entries=min_project_entries,
        min_project_bullets=min_project_bullets,
    )

    def _trim_static_one_step(selection: dict[str, Any]) -> tuple[dict[str, Any], bool]:
        """Trim by lowest relevance score when available, else by position."""
        if selection.get("content_composites"):
            updated, event = trim_selection_by_lowest_score(selection, policy=trim_policy)
            return updated, event is not None
        if not selection_trim_exhausted(
            selection,
            min_project_entries=min_project_entries,
            min_project_bullets=min_project_bullets,
        ):
            return (
                trim_selection_one_step(
                    selection,
                    min_project_entries=min_project_entries,
                    min_project_bullets=min_project_bullets,
                ),
                True,
            )
        return selection, False

    if auto_fill_page_budget and current_selection.get("content_composites"):
        for _ in range(50):
            expanded, event = expand_selection_by_highest_score(
                current_selection, background_data
            )
            if event is None:
                break
            _trial_docx, trial_pdf, trial_pages = _render_and_measure(expanded, current)
            if trial_pdf is None or trial_pages is None or trial_pages > max_pages:
                break
            current_selection = expanded
            docx_bytes = _trial_docx
            pdf_bytes = trial_pdf
            page_count = trial_pages

    for _ in range(30):
        if page_count <= max_pages:
            break
        current_selection, progressed = _trim_static_one_step(current_selection)
        if not progressed and include_summary:
            trimmed = trim_summary_one_step(current, sections=sections)
            if trimmed != current:
                current = trimmed
                progressed = True
        if not progressed:
            trim_stalled = page_count > max_pages
            break

        docx_bytes, pdf_bytes, page_count = _render_and_measure(current_selection, current)
        if pdf_bytes is None:
            break

    if page_count is not None and page_count > max_pages:
        trim_stalled = True

    diagnostics = _page_fit_diagnostics(
        background_data,
        current,
        current_selection,
        page_count=page_count,
        max_pages=max_pages,
        max_chars=max_chars,
        include_summary=include_summary,
        include_skills=include_skills,
        trim_stalled=trim_stalled,
        overflow_warning_min_composite=overflow_warning_min_composite,
        render_page_count=_render_and_measure,
        ai_output_for_render=current,
    )
    return current, current_selection, docx_bytes, pdf_bytes, page_count, diagnostics


def _page_fit_diagnostics(
    background_data: dict[str, Any],
    ai_output: dict[str, Any],
    content_selection: dict[str, Any],
    *,
    page_count: int | None,
    max_pages: int,
    max_chars: int | None,
    include_summary: bool,
    include_skills: bool,
    trim_stalled: bool,
    overflow_warning_min_composite: float = DEFAULT_OVERFLOW_WARNING_MIN_COMPOSITE,
    render_page_count: Any | None = None,
    ai_output_for_render: dict[str, Any] | None = None,
) -> dict[str, Any]:
    stats = static_content_stats(background_data, content_selection)
    fitted_chars = ai_output_char_count(
        ai_output["summary"],
        ai_output["skill_categories"],
        include_summary=include_summary,
        include_skills=include_skills,
    )

    overflow_warning = False
    overflow_message: str | None = None
    omitted_high_quality_count = 0

    if (
        page_count is not None
        and render_page_count is not None
        and ai_output_for_render is not None
        and content_selection.get("content_composites")
    ):
        omitted = high_quality_omitted_items(
            content_selection,
            background_data,
            min_composite=overflow_warning_min_composite,
        )
        omitted_high_quality_count = len(omitted)
        if omitted:
            trial_selection = selection_with_all_items(content_selection, omitted)
            _docx, trial_pdf, trial_pages = render_page_count(
                trial_selection, ai_output_for_render
            )
            if trial_pages is not None and trial_pages >= max_pages + 1:
                overflow_warning = True
                overflow_message = (
                    f"Enough strong job-relevant content remains to fill a full additional "
                    f"page ({omitted_high_quality_count} high-scoring items omitted). "
                    f"Consider setting max_resume_pages = {max_pages + 1}."
                )

    return {
        **stats,
        "page_count": page_count,
        "max_pages": max_pages,
        "ai_char_count": fitted_chars,
        "ai_char_budget": max_chars,
        "trim_stalled": trim_stalled,
        "trim_log": list(content_selection.get("trim_log", [])),
        "expand_log": list(content_selection.get("expand_log", [])),
        "overflow_warning": overflow_warning,
        "overflow_message": overflow_message,
        "omitted_high_quality_count": omitted_high_quality_count,
    }


def build_resume_artifacts(
    background_data: dict[str, Any],
    ai_output: dict[str, Any],
    *,
    max_pages: int = 2,
    max_chars: int | None = None,
    sections: list[str] | None = None,
    content_selection: dict[str, Any] | None = None,
    max_certifications: int | None = None,
    min_project_entries: int = DEFAULT_MIN_PROJECT_ENTRIES,
    min_project_bullets: int = DEFAULT_MIN_PROJECT_BULLETS,
    auto_fill_page_budget: bool = True,
    overflow_warning_min_composite: float = DEFAULT_OVERFLOW_WARNING_MIN_COMPOSITE,
) -> dict[str, Any]:
    """Build DOCX and derived preview/export artifacts from background + AI output."""
    fitted_output, fitted_selection, docx_bytes, pdf_bytes, page_count, diagnostics = (
        fit_resume_to_page_budget(
            background_data,
            ai_output,
            max_pages=max_pages,
            max_chars=max_chars,
            sections=sections,
            content_selection=content_selection,
            max_certifications=max_certifications,
            min_project_entries=min_project_entries,
            min_project_bullets=min_project_bullets,
            auto_fill_page_budget=auto_fill_page_budget,
            overflow_warning_min_composite=overflow_warning_min_composite,
        )
    )
    return {
        "ai_output": fitted_output,
        "content_selection": fitted_selection,
        "docx_bytes": docx_bytes,
        "html": docx_to_html(docx_bytes),
        "pdf_bytes": pdf_bytes,
        "page_count": page_count,
        "diagnostics": diagnostics,
    }
